"""
Extractor para archivos Word DOCX usando python-docx.
"""

from docx import Document
from pathlib import Path
from typing import Dict, Any, List
import logging
import re

from .base_extractor import BaseExtractor

logger = logging.getLogger(__name__)

class DOCXExtractor(BaseExtractor):
    """
    Extractor para archivos Word DOCX.
    Extrae texto, preserva estructura de párrafos y tablas.
    """
    
    def __init__(self):
        super().__init__()
        self.supported_formats = [".docx"]
    
    async def extract(self, file_path: Path) -> Dict[str, Any]:
        """
        Extrae contenido de archivo DOCX.
        
        Args:
            file_path: Ruta al archivo .docx
            
        Returns:
            Dict con contenido y metadata
        """
        try:
            # Validar archivo
            if not self.validate_file(file_path):
                raise ValueError(f"Archivo DOCX inválido: {file_path}")
            
            # Abrir documento
            doc = Document(file_path)
            
            # Extraer contenido
            paragraphs = []
            tables = []
            full_content = ""
            
            # Procesar párrafos
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    paragraphs.append({
                        "index": para_idx,
                        "text": text,
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    })
                    full_content += text + "\n\n"
            
            # Procesar tablas
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table(table, table_idx)
                tables.append(table_data)
                full_content += f"\n[TABLA {table_idx + 1}]\n{table_data['text']}\n\n"
            
            # Limpiar contenido completo
            full_content = self._clean_text(full_content)
            
            # Extraer propiedades del documento
            core_properties = doc.core_properties
            
            # Crear metadata
            metadata = self.create_metadata(file_path, {
                "paragraph_count": len([p for p in paragraphs if p["text"]]),
                "table_count": len(tables),
                "char_count": len(full_content),
                "estimated_tokens": self.estimate_tokens(full_content),
                "document_properties": {
                    "title": core_properties.title or "",
                    "author": core_properties.author or "",
                    "subject": core_properties.subject or "",
                    "created": str(core_properties.created) if core_properties.created else "",
                    "modified": str(core_properties.modified) if core_properties.modified else ""
                }
            })
            
            # Crear chunks
            chunks = self._create_chunks(paragraphs, tables)
            
            return {
                "content": full_content,
                "metadata": metadata,
                "chunks": chunks
            }
            
        except Exception as e:
            logger.error(f"Error extrayendo DOCX {file_path}: {e}")
            raise
    
    def _extract_table(self, table, table_idx: int) -> Dict[str, Any]:
        """
        Extrae contenido de una tabla.
        
        Args:
            table: Objeto tabla de python-docx
            table_idx: Índice de la tabla
            
        Returns:
            Dict con datos de la tabla
        """
        rows_data = []
        table_text = ""
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            
            rows_data.append(row_data)
            table_text += " | ".join(row_data) + "\n"
        
        return {
            "index": table_idx,
            "rows": rows_data,
            "text": table_text.strip(),
            "row_count": len(rows_data),
            "col_count": len(rows_data[0]) if rows_data else 0
        }
    
    def _clean_text(self, text: str) -> str:
        """
        Limpia texto extraído de DOCX.
        
        Args:
            text: Texto crudo
            
        Returns:
            str: Texto limpio
        """
        if not text:
            return ""
        
        # Normalizar espacios y saltos de línea
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remover líneas vacías al inicio y final
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line or lines[lines.index(line)-1:lines.index(line)+2].count(''))
        
        return text.strip()
    
    def _create_chunks(self, paragraphs: List[Dict], tables: List[Dict]) -> List[Dict]:
        """
        Crea chunks basados en párrafos y tablas.
        
        Args:
            paragraphs: Lista de párrafos
            tables: Lista de tablas
            
        Returns:
            List de chunks
        """
        chunks = []
        chunk_id = 0
        
        # Chunks de párrafos
        for paragraph in paragraphs:
            text = paragraph["text"]
            if len(text) > 50:  # Solo párrafos significativos
                chunk_id += 1
                
                chunks.append({
                    "content": text,
                    "chunk_id": f"docx_para_{chunk_id}",
                    "metadata": {
                        "type": "paragraph",
                        "paragraph_index": paragraph["index"],
                        "style": paragraph["style"],
                        "length": len(text),
                        "tokens": self.estimate_tokens(text)
                    }
                })
        
        # Chunks de tablas
        for table in tables:
            if table["text"] and len(table["text"]) > 20:
                chunk_id += 1
                
                chunks.append({
                    "content": table["text"],
                    "chunk_id": f"docx_table_{chunk_id}",
                    "metadata": {
                        "type": "table",
                        "table_index": table["index"],
                        "row_count": table["row_count"],
                        "col_count": table["col_count"],
                        "length": len(table["text"]),
                        "tokens": self.estimate_tokens(table["text"])
                    }
                })
        
        return chunks